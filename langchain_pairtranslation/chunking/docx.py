from functools import cached_property, cache
import re
from typing import Annotated, Iterator
import annotated_types
import docx.document
import logging
from re import Pattern
from pydantic import BaseModel, Field, PrivateAttr, PositiveInt
from langchain_pairtranslation.chunking.base import BaseDocumentChunks
from langchain_pairtranslation.config import SplitPattern, SplitBehavior
from langchain_pairtranslation.rowinfo import RowCol, RowMetadata
from langchain_pairtranslation.translation_state import TranslationState
LOG = logging.getLogger('app')

import langchain_pairtranslation.utils.string_extensions as string

class DocxChunks(BaseDocumentChunks):
    """
    A series of DocxChunks drawing from a docx.Document object.
    """
    @cached_property
    def src(self) -> docx.document.Document:
        return docx.Document(self.doc_path)

    @classmethod
    def from_docx(cls, doc_path: str) -> 'DocxChunks':
        open_document = docx.Document(doc_path)
        raw_strings = [paragraph.text for paragraph in open_document.paragraphs]
        chunks = BaseDocumentChunks.from_strings(raw_strings)
        instance = cls(doc_path=doc_path, collection=chunks)
        instance.src = open_document
        return instance
    
    def _get_row(self, index: int) -> str:
        return self.src.paragraphs[index].text
    
    def split_pass(self, max_length: int, splitters: list[SplitPattern]) -> None:
        super().split_pass(self._get_row, max_length, splitters)
    
    def to_row_overrides_iter(self) -> Iterator[tuple[RowMetadata, str]]:
        for chunk in self.collection:
            for row in chunk.to_rows(self._get_row):
                yield row

    def to_row_exports_iter(self) -> Iterator[str]:
        return super().to_row_exports_iter(self._get_row, len(self.src.paragraphs))