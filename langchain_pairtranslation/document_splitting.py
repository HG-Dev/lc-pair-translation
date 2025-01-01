from enum import Enum
from functools import cached_property
from typing import Annotated, Callable, Iterator
import annotated_types
import docx
import docx.document
import logging
from frozenlist import FrozenList
from pydantic import BaseModel, Field, PositiveInt
import re
from functools import total_ordering
UnsignedInt = Annotated[int, annotated_types.Ge(0)]

import langchain_pairtranslation.string_extensions as string

log = logging.getLogger("app")

@total_ordering #TODO: Don't use total_ordering, it's reportedly slower than implementing the rest
class RowCol(BaseModel):
    """
    ### Summary
    A row and column pair representing a position in a document.
    ### Attributes
        1. row : PositiveInt
        2. column : PositiveInt
    """
    row: PositiveInt = Field(default=0)
    column: PositiveInt = Field(default=0)

    def __eq__(self, other):
        for left, right in zip(iter(self), iter(other)):
            if left != right:
                return False
        return True
    
    def __lt__(self, other):
        for left, right in zip(iter(self), iter(other)):
            if left < right:
                return True
            elif right > left:
                return False
        return False

    @staticmethod
    def zero():
        return RowCol()
    
    @classmethod
    def from_row_start(cls, row: PositiveInt):
        return cls(row=row, column=1)
    
    @classmethod
    def create(cls, row: PositiveInt, column: PositiveInt):
        return cls(row=row, column=column)
    
    def dot_str(self) -> str:
        return f'{self.row}.{self.column}'
    
    def to_tuple(self, zero_indexed = False) -> tuple[int, int]:
        if zero_indexed:
            return self.row - 1, self.column
        return self.row, self.column

    def __str__(self):
        return f"(Row {self.row}, Col {self.column})"
    
    def __iter__(self):
        yield self.row
        yield self.column

    def is_within_bounds(self, max_row, max_column):
        """Check if the coordinate is within the given bounds."""
        return 0 <= self.row < max_row and 0 <= self.column < max_column


class DocxChunk(BaseModel):
    """Document Chunks are [symantically] cohesive portions of a document sent to an LLM to be translated.

    Chunks can be any length, depending on what's used to break the document into chunks.
    This implies a need to direct translation services to generate output one portion at a time for each individual chunk.
    """
    class TranslationState(Enum):
        UNTRANSLATED = 0
        LLM_TRANSLATED = 1
        USER_TRANSLATED = 2
        USER_EDITED = 3

    start_: RowCol = Field(
        description="The row and column at which this chunk begins (inclusive, one indexed). "
                        + "The presence of zeroes implies the value has not yet been set.",
        default = RowCol.create(1,1),
        alias='start'
    )
    end_: RowCol = Field(
        description="The row and column at which this chunk terminates (inclusive, one indexed). "
                        + "The presence of zeroes implies the value has not yet been set.",
        default = RowCol.create(1,1),
        alias='end'
    )
    state_: TranslationState = Field(
        description="Whether this chunk is translated, and if so, by who.",
        default=TranslationState.UNTRANSLATED,
        alias='state'
    )

    @property
    def is_src_set(self):
        return self.__src is not None

    def set_src(self, source: docx.document.Document):
        self.__src = source

    def expand_until(self, divider: re.Pattern, inclusive: bool) -> RowCol:
        """ Expand this chunk until a paragraph that matches or contains the given divider is found.
        ### Parameters
        1. divider : str
                - The string or regex pattern that identifies where the chunk should stop expanding.
        2. inclusive : bool
                - Whether or not the chunk should in corporate the divider into its span.
        ### Returns
        The RowCol from which the next chunk should start from.
        """
        if (self.__src is None):
            print("SRC = NONE")
            self.end_.row += 1
            return self.end_

        assert(self.start_.row > 0 and self.end_.row > 0)
        
        paragraphs = self.__src.paragraphs
        source_row_count = len(paragraphs)
        # Extend the end of the span until divider is found
        while self.end_.row <= source_row_count:
            paragraph = paragraphs[self.end_.row - 1].text
            paragraph_length = len(paragraph)

            previous_end = self.end_

            # Start by assuming end should be at the end of this line
            self.end_.column = paragraph_length

            found = divider.match(paragraph)
            if found:
                pos, length = found.span()
                print(f"Found pattern {divider.pattern} on row {self.end_.row} @ {pos}->{length}")
                matched_entire_line = length == paragraph_length
                
                if inclusive:
                    # Push pos up to increase end
                    pos += length
                
                # Divider is on the start of a line
                if pos == 0:
                    # end still has len of previous line; move it back there
                    self.end_ = previous_end
                    if matched_entire_line: # Skip over next line
                        return RowCol.from_row_start(self.end_.row + 2)
                
                raise NotImplementedError()
                # if self.end.column >= paragraph_length:
                #     return RowCol.from_row_start(self.end.row + 1)

                # Next chunk will begin within the same paragraph(!)
                return RowCol.create(self.end_.row, self.end_.column+1)
            # Expansion can continue
            self.end_.row += 1
        
        # Reached the end of the document
        self.end_ = RowCol.create(source_row_count, len(paragraphs[-1].text))
        # This 'next' will be outside the range of the document
        return RowCol.from_row_start(self.end_.row + 1) 

    @property
    def start(self) -> RowCol:
        return self.start_
    
    @property
    def end(self) -> RowCol:
        return self.end_
    
    @property
    def ref_paragraph_count(self) -> int:
        return (self.end.row - self.start.row) + 1

    # TODO: Columns
    @property
    def src_text_substrings(self) -> Iterator[str]:
        if self.__src is None or self.start.row < 1:
            return
        source = self.__src
        if self.start.row > len(self.__src.paragraphs):
            raise ValueError(f"Chunk starts on {self.start} outside of document space ({len(self.__src.paragraphs)}). The chunk should not have been created.")
        if self.end.row == self.start.row:
            yield self.__src.paragraphs[self.start.row - 1].text[self.start.column-1:self.end.column]
        else:
            # Yield entire starting line from src_start.column
            yield self.__src.paragraphs[self.start.row - 1].text[self.start.column-1:]
            for idx in range(self.start.row, self.end.row - 2):
                # Yield all complete lines between start and end
                yield source.paragraphs[idx].text
            # Yield ending line ending at src_end.column
            yield self.__src.paragraphs[self.end.row - 1].text[:self.end.column]

    @property
    def src_text(self) -> str:
        if self.__src is None:
            return ""
        output = self.__src.paragraphs[self.start.row - 1].text[self.start.column-1:]
        return '\n'.join([substring for substring in self.src_text_substrings])

    @cached_property
    def src_text_length(self, include_line_breaks=False):
        total = sum([len(substring) for substring in self.src_text_substrings])
        lines = self.ref_paragraph_count
        if not include_line_breaks:
            return total
        else:
            return total + lines
        
    @property
    def state(self) -> TranslationState:
        return self.state_

    # def __init__(self, should_finalize_chunk: Callable[[str, int], bool]):
    #     self._src_paragraphs: FrozenList[str] = FrozenList() # Paragraphs from source document
    #     self._src_length = 0
    #     self._src_token_count = 0
    #     self._check_paragraph_for_divider: Callable[[str, int], bool] = should_finalize_chunk
    #     self._out_paragraphs: List[str] = []                 # Translated paragraphs
    #     self._start = RowCol.zero()               # Row, column where src paragraphs begin
    #     self._end = RowCol.zero()                 # Row, column where src paragraphs end
    #     self._state: DocumentChunk.TranslationState = DocumentChunk.TranslationState.UNTRANSLATED

    # def append_src_paragraph(self, p: str, p_idx: int):
    #     ''' Add or initialize src data with next paragraph '''
    #     if self._start.row == 0:
    #         self._start = RowCol.create(p_idx, 0)
    #         logger.debug("Src paragraphs set to start at {}".format(self._start))
    #         self._src_length = 0

    #     self._src_paragraphs.append(p)

    #     p_length = len(p)
    #     prev_end = self._end
    #     self._end = RowCol.create(p_idx, p_length)
    #     self._src_length += p_length
    #     logger.debug("Added src paragraph of length {}. End advanced from {} to {}".format(p_length, prev_end, self._end))

    #     # Clear token count cache (if necessary in the first place)
    #     self.__dict__.pop('src_token_count', None)

    #     if self._check_paragraph_for_divider(p, self._src_length):
    #         self.finalize_src()

    # def finalize_src(self):
    #     ''' Remove excess whitespace from the beginning and ends, then freeze internal frozen list of source paragraphs '''
    #     all_removed = ""

    #     while string.is_empty_or_whitespace(self._src_paragraphs[-1]):
    #         all_removed += self._src_paragraphs.pop()
    #         self._end = RowCol.create(self._end.row - 1, len(self._src_paragraphs[-1]))
    #     while string.is_empty_or_whitespace(self._src_paragraphs[0]):
    #         all_removed += self._src_paragraphs.pop(0)
    #         self._start = RowCol.create(self._start.row + 1, 0)
        
    #     self._src_length -= len(all_removed)
    #     self._src_paragraphs.freeze()

    # def set_llm_translation(self, text: str):
    #     self._out_paragraphs = text.split('\n')
    #     self._state = DocumentChunk.TranslationState.LLM_TRANSLATED

    # @property
    # def src_length(self) -> int:
    #     return self._src_length
    
    # @property
    # def src_paragraph_count(self) -> int:
    #     return len(self._src_paragraphs)
    
    # @property
    # def state(self) -> TranslationState:
    #     return self._state
    
    # @cached_property
    # def src_token_count(self) -> int:
    #     raise NotImplementedError("Use src_length for now, which identifies the 'worst-case' amount of tokens.")

    # @property
    # def src_is_finalized(self) -> bool:
    #     return self._src_paragraphs.frozen
    
    # @property
    # def src_is_empty(self) -> bool:
    #     if self._src_length == 0:
    #         return True
    #     return all(string.is_empty_or_whitespace(p) for p in self._src_paragraphs)

    # @property
    # def src_start(self) -> RowCol:
    #     return self._start
    
    # @property
    # def src_end(self) -> RowCol:
    #     return self._end

    # @property
    # def src_paragraphs(self) -> Iterator[str]:
    #     return self._src_paragraphs.__iter__()

class DocxChunks(BaseModel):
    """
    A series of DocxChunks drawing from a docx.Document object.
    """
    # Public fields
    doc_path: str = Field(
        description="The filepath of the source docx.",
        frozen=True
    )
    chunks_: list[DocxChunk] = Field(
        description="The internal collection of document chunks.",
        alias='chunks',
        frozen=True,
        default=[]
    )
    next_start_: RowCol = Field(
        description="The position from which the next created chunk should be created",
        default=RowCol.zero()
    )

    @staticmethod
    def load_from_document_path(doc_path: str):
        doc = docx.Document(doc_path)
        
        return DocxChunks(doc_path=doc_path)

    @cached_property
    def _doc(self) -> docx.document.Document:
        ''' The source document from which chunks are derived '''
        open_doc = docx.Document("resources/test.docx")
        print("Opened document with {} pages and {} paragraphs".format(len(open_doc.sections), len(open_doc.paragraphs)))
        return open_doc
    
    @cached_property
    def _doc_paragraph_count(self) -> UnsignedInt:
        return len(self._doc.paragraphs)

    @property
    def collection(self):
        return self.chunks_
    
    def __iter__(self):
        for chunk in self.collection:
            yield chunk
    
    def split_document(self):
        while self.next_start_.row <= self._doc_paragraph_count:
            prev_value = self.next_start_
            created, self.next_start_ = self.create_chunk_with_start(self.next_start_, self._doc)
            try:
                assert(self.next_start_.row > self._doc_paragraph_count or prev_value < self.next_start_)
            except AssertionError:
                row_idx = prev_value.row - 1
                print(f"Failed to advance from {prev_value} :\n{self._doc.paragraphs[row_idx-1].text if row_idx > 0 else "DOC START"}\n---->{self._doc.paragraphs[row_idx].text}")
                break
            assert(created.start.row <= self._doc_paragraph_count)
            self.collection.append(created)
            
    @staticmethod #source_getter: Callable[..., docx.document.Document]
    def create_chunk_with_start(start: RowCol, source) -> tuple[DocxChunk, RowCol]:
        chunk = DocxChunk(start=start, end=start)
        print("Created chunk at ", start)
        chunk.set_src(source)
        assert(source is not None)
        return (chunk, chunk.expand_until(re.compile(r'^\s*$'), inclusive=False))
        

