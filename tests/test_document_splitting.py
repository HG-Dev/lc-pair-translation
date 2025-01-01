"""
Tests the division of various types of documents into translatable chunks.
"""
import unittest
import docx

from langchain_pairtranslation import document_splitting
from tests import pytest_log

class DocumentSplitTests(unittest.TestCase):
        
    def test_wagahai(self):
        """
        Creates TranslatableChunks from the text of a Japanese novel.
        Paragraphs are sequential without any spaces between them.
        Includes ruby furigana text that should be stripped before translation.
        """
        pytest_log.info("Opening 'wagahai.docx'...")
        doc = docx.Document("tests/resources/wagahai.docx")
        pytest_log.info(f"Opened 'wagahai.docx'; {len(doc.paragraphs)} paragraphs found.")
        self.assertGreater(len(doc.paragraphs), 0, "No paragraphs found in 'wagahai.docx'.")