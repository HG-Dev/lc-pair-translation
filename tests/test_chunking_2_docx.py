"""
Tests the division of various types of docx documents into translatable chunks.
"""
import unittest
import docx
import re
from langchain_pairtranslation.chunking.docx import DocxChunks
from langchain_pairtranslation.config import Config, SplitBehavior, SplitPattern
from tests import pytest_log

class DocxChunkingTests(unittest.TestCase):

    def test_gettysburg_balanced_sentence_splitting(self):
        """
        Confirm that the Gettysburg Address is split into chunks of balanced length.
        The Gettysburg Address has a total of 1458 characters, but with a max length of 360,
        the final sentence will exceed the max length.
        """
        chunks = DocxChunks.from_docx("tests/resources/gettysburg.docx")
        assert(len(chunks) == 1)
        
        splitter = SplitPattern(name="End of sentence", pattern_src=r'([.!?。！？…])[^A-z0-9"]')
        chunks.split_pass(360, [splitter])
        self.assertEqual(len(chunks), 4)
        first_three_avg = sum([len(chunk) for chunk in chunks[:3]]) / 3
        self.assertTrue((360 - first_three_avg) < 60, 
                        f"First three chunks are not balanced in length: {', '.join([str(len(chunk)) for chunk in chunks])}")

    def test_wagahai(self):
        """
        Creates translateable chunks from the text of a Japanese novel.
        Paragraphs are sequential without any spaces between them.
        Includes ruby furigana text that should be stripped before translation.
        """
        chunks = DocxChunks.from_docx("tests/resources/wagahai.docx")
        doc = chunks.src
        pytest_log.info(f"Created chunks from 'wagahai.docx'; {len(chunks)} out of {len(doc.paragraphs)} paragraphs converted into chunks.")
        self.assertEqual(len(chunks), len(doc.paragraphs))
        for src_paragraph, chunk in zip(doc.paragraphs, chunks):
            self.assertEqual(len(src_paragraph.text.strip()), len(chunk), "Chunk length does not match source paragraph length.")

        config = Config().chunking
        chunks.merge_pass(config.min_length)
        chunks.split_pass(config.max_length, config.split_patterns)
        pytest_log.debug("Refined chunks: %d", len(chunks))
        for chunk in chunks:
            pytest_log.debug(f"({str(chunk)}), len={len(chunk)}: {chunk.to_text(chunks._get_row)}")

        untranslated_exported_rows = [row for row in chunks.to_row_exports_iter()]
        for (in_row, out_row) in zip(doc.paragraphs, untranslated_exported_rows):
            self.assertEqual(in_row.text, out_row, "Untranslated row does not match source row.")

    @unittest.skip("Needs source document")
    def test_art_event(self):
        """
        Creates translateable chunks from the text of a tourism article.
        """
        chunks = DocxChunks.from_docx("tests/resources/art_event.docx")
        doc = chunks.src
        pytest_log.info(f"Created chunks from 'art_event.docx'; {len(chunks)} out of {len(doc.paragraphs)} paragraphs converted into chunks.")

        config = Config().chunking
        chunks.merge_pass(config.min_length)
        chunks.split_pass(config.max_length, config.split_patterns)
        pytest_log.debug("Refined chunks: %d", len(chunks))
        for chunk in chunks:
            text = chunk.to_text(chunks._get_row)
            pytest_log.debug(f"({str(chunk)}), len={len(chunk)}/{len(text)} multiline={'\n' in text}: {chunk.to_text(chunks._get_row)}")

        untranslated_exported_rows = [row for row in chunks.to_row_exports_iter()]
        for (in_row, out_row) in zip(doc.paragraphs, untranslated_exported_rows):
            self.assertEqual(in_row.text, out_row, "Untranslated row does not match source row.")


        
            