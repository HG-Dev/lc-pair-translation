from typing import Union, Tuple, Optional, Callable
from tkinter.constants import *
from customtkinter.windows.widgets.theme import ThemeManager
from langchain_core.prompts import ChatPromptTemplate, SystemMessagePromptTemplate, HumanMessagePromptTemplate, AIMessagePromptTemplate
import customtkinter as ctk
import docx
import logging
import sys
import asyncio

from langchain_pairtranslation.customtkinter_extensions import *
from langchain_pairtranslation.document_analysis import *
from langchain_pairtranslation.llms import BaseLLMProvider
import langchain_pairtranslation.string_extensions as string
#from langchain_pairtranslation._internal import MyLLMProvider
from langchain_pairtranslation.app_config import Config
from langchain_pairtranslation.document_splitting import DocxChunks, DocxChunk

CHUNK_DIVIDER="\n"

logger = logging.getLogger(__name__)
logger.setLevel(logging.DEBUG)
handler = logging.StreamHandler(sys.stdout)
handler.setFormatter(logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s'))
handler.setLevel(logging.DEBUG)
logger.addHandler(handler)

class TranslationController:
    def __init__(self, master: ctk.CTkFrame, config: Config, chunks: List[DocxChunk]):
        master.grid_columnconfigure(0, weight=1)
        master.grid_columnconfigure(1, weight=0)
        master.grid_columnconfigure(2, weight=1)
        master.grid_rowconfigure(0, weight=1, minsize=64-10)
        master.grid_rowconfigure(1, weight=1, minsize=10)
        self._config = config
        self.main_button_txt = ctk.StringVar(master, value="Translate Next Chunk")
        self.main_button = ctk.CTkButton(master, width=256, height=32)
        self.main_button.configure(command=self.start_translation, textvariable=self.main_button_txt)
        self.progressbar = ctk.CTkProgressBar(master, mode="indeterminate")
        self.progressbar_placeholder = ctk.CTkFrame(master, bg_color="transparent", fg_color="transparent", height=self.progressbar._current_height)
        self.main_button.grid(column=1, row=0, pady=2, sticky=S)
        self.progressbar_placeholder.grid(column=1, row=1, pady=2, sticky="SWE")
        self._chunks = chunks

    def _show_progressbar(self, value: bool):
        if value:
            self.progressbar_placeholder.grid_forget()
            self.progressbar.grid(column=1, row=1, pady=2, sticky="SWE")
            self.progressbar.start()
        else:
            self.progressbar.stop()
            self.progressbar.grid_forget()
            self.progressbar_placeholder.grid(column=1, row=1, pady=2, sticky="SWE")

    def start_translation(self, **kwargs):
        self.main_button.configure(state=DISABLED)
        for key, value in kwargs.items():
            print(key, value)
        self._show_progressbar(True)
        # Get next untranslated chunk
        for chunk in self._chunks:
            if chunk.state == DocxChunk.TranslationState.UNTRANSLATED:
                logger.info("Found untranslated chunk: {}".format(string.take_multi(chunk.src_text_substrings, 24, "...").replace('\n', '¶')))
                asyncio.run(self._start_translation_internal(chunk, config))
                
                return
        logger.warning("All chunks are translated!")
        #asyncio.run()

    @staticmethod
    async def _start_translation_internal(chunk: DocxChunk, config: Config):
        # Prepare LLM and prompt
        #llm = await MyLLMProvider().create_translator()
        translator_kwargs = config.translator_format_kwargs
        sysprompt = config.translator_base_system_prompt
        # terms = [
        #     DocumentTerm(source="in", translations=["out"], description="DO NOT SAVE THIS"),
        # ]
        #print(string.wrap_with_xml_tag(str(config.terms), "TERMS"))

        sysP = SystemMessagePromptTemplate.from_template("System prompt {test}")
        exHuman = HumanMessagePromptTemplate.from_template("Hello, world!")
        exAI = AIMessagePromptTemplate.from_template("HELLOTE W0RLD!")
        promptTemplate = ChatPromptTemplate.from_messages([sysP, exHuman, exAI])
        print(promptTemplate.format_prompt(test="test successful").to_string())
    

def create_root() -> ctk.CTk:
    root = ctk.CTk()
    root.title("LangChain Paired Translation")
    root.geometry('1280x960')
    #root.iconbitmap('resoures/something.ico')
    ctk.FontManager.load_font('resources/fonts/RobotoMono-Regular.ttf')
    ctk.set_appearance_mode("dark")
    ctk.set_default_color_theme("dark-blue")
    return root

if __name__ == '__main__':
    config = Config("config/settings.ini")
    root = create_root()
    status_frame = ctk.CTkFrame(root, height=24)
    status_frame.pack(side=BOTTOM, fill=X)

    ctrl_frame = ctk.CTkFrame(root, height=64, bg_color=ThemeManager.theme["CTkScrollableFrame"]["label_fg_color"])
    ctrl_frame.pack(side=TOP, fill=X)

    text_frame = ChunkedTranslationFrame(root, bg_color="black", fg_color="black")
    text_frame.pack(side=TOP, fill=BOTH, expand=TRUE)

    #document = docx.Document('resources/src2.docx')
    #logger.info(f"Opened document with {len(document.paragraphs)} paragraphs: {document.core_properties.title}" )
    # chunk_text_buffer = ""
    # chunk_start = RowCol.create(1, 0)
    # last_significant_end = RowCol.create(1, 0)
    # chunk_end = RowCol.create(1, 0)
    # divider_regex = config.current_chunk_divider_regex
    # logger.debug("Chunk divider pattern received: {}".format(divider_regex))
    # min_chunk_length = config.min_chunk_length

    # def should_finalize_chunk(p: str, chunk_size: int):
    #     return chunk_size > min_chunk_length and divider_regex.search(p)

    chunks = DocxChunks(doc_path="")

    chunks.split_document()

    # current_chunk = (should_finalize_chunk)

    # for p_idx, paragraph in enumerate(document.paragraphs):

    #     if current_chunk.src_is_finalized:
    #         logger.debug(f"Chunk finalized: {current_chunk.src_start} {current_chunk.src_end}")
    #         text_frame.add_chunk(current_chunk)
    #         current_chunk = DocumentChunk(should_finalize_chunk)

    #     current_chunk.append_src_paragraph(paragraph.text, p_idx + 1)
    #     #chunk_text_buffer += paragraph.text.replace('…', '...') + '\n'
    
    # # Finalize and add chunk at the end of the document-- if anything noteworthy exists inside it
    # current_chunk.finalize_src()
    # if not current_chunk.src_is_empty:
    #     text_frame.add_chunk(current_chunk)
    for chunk in chunks:
        text_frame.add_chunk(chunk)

    ctrls = TranslationController(ctrl_frame, config, text_frame._chunks)

    root.mainloop()

